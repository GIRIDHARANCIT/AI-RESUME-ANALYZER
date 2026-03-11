"""Convert resume text to different formats (PDF, DOCX, TXT)."""
from pathlib import Path
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


def _escape_html(s: str) -> str:
    """Escape HTML special characters for PDF generation."""
    return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def text_to_pdf(text: str, output_path: Path) -> Path:
    """Convert plain text resume to PDF."""
    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=letter,
        rightMargin=0.75 * inch,
        leftMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
    )
    styles = getSampleStyleSheet()
    elements = []
    for block in text.split("\n\n"):
        lines = block.split("\n")
        safe_lines = [_escape_html(l) for l in lines if l.strip()]
        if not safe_lines:
            elements.append(Spacer(1, 0.1 * inch))
        else:
            para = "<br/>".join(safe_lines)
            elements.append(Paragraph(para, styles["Normal"]))
            elements.append(Spacer(1, 0.1 * inch))
    doc.build(elements)
    return output_path


def text_to_docx(text: str, output_path: Path) -> Path:
    """Convert plain text resume to DOCX."""
    doc = Document()
    
    for block in text.split("\n\n"):
        lines = [l for l in block.split("\n") if l.strip()]
        if not lines:
            continue
            
        # Check if this looks like a section header (all caps or bold indicators)
        first_line = lines[0]
        if first_line.isupper() and len(first_line) < 50:
            # Add as heading
            p = doc.add_heading(first_line, level=1)
            p.style = 'Heading 1'
            # Add remaining lines
            for line in lines[1:]:
                if line.strip():
                    doc.add_paragraph(line)
        else:
            # Add as regular paragraphs
            for line in lines:
                if line.strip():
                    doc.add_paragraph(line)
        
        # Add spacing between blocks
        doc.add_paragraph()
    
    doc.save(str(output_path))
    return output_path


def text_to_txt(text: str, output_path: Path) -> Path:
    """Save plain text resume as TXT file."""
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(text)
    return output_path
