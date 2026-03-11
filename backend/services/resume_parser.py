"""Resume text extraction from multiple file formats."""
import re
from pathlib import Path
from typing import Optional

from pdfminer.high_level import extract_text as pdf_extract
from docx import Document


def _clean_text(text: str) -> str:
    """Clean and normalize extracted text."""
    if not text:
        return ""
    # Replace multiple whitespace with single space
    text = re.sub(r"\s+", " ", text)
    # Remove control characters
    text = "".join(c for c in text if c.isprintable() or c in "\n\t")
    return text.strip()


def extract_from_pdf(filepath: Path) -> str:
    """Extract text from PDF using pdfminer."""
    try:
        text = pdf_extract(str(filepath))
        return _clean_text(text or "")
    except Exception as e:
        raise ValueError(f"Failed to parse PDF: {str(e)}")


def extract_from_docx(filepath: Path) -> str:
    """Extract text from DOCX using python-docx."""
    try:
        doc = Document(filepath)
        paragraphs = [p.text for p in doc.paragraphs]
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    tables_text.append(cell.text)
        full_text = "\n".join(paragraphs) + "\n" + "\n".join(tables_text)
        return _clean_text(full_text)
    except Exception as e:
        raise ValueError(f"Failed to parse DOCX: {str(e)}")


def extract_from_text_file(filepath: Path) -> str:
    """Extract text from plain text files (.txt, .rtf, .log, etc)."""
    try:
        with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
        return _clean_text(text)
    except Exception as e:
        raise ValueError(f"Failed to parse text file: {str(e)}")


def extract_text(filepath: Path) -> str:
    """
    Extract text from resume file (supports all common formats).
    Returns cleaned, tokenizable text.
    
    Supported formats:
    - PDF (.pdf)
    - Word (.docx, .doc)
    - Text (.txt, .rtf, .log, .odt, etc.)
    
    Falls back to plain text extraction for unknown formats.
    """
    suffix = filepath.suffix.lower()
    
    try:
        if suffix == ".pdf":
            return extract_from_pdf(filepath)
        elif suffix == ".docx":
            return extract_from_docx(filepath)
        else:
            # Try to extract as text for any other file type
            return extract_from_text_file(filepath)
    except Exception as e:
        # Final fallback: try text extraction
        try:
            return extract_from_text_file(filepath)
        except Exception:
            raise ValueError(f"Unable to extract text from {suffix} file: {str(e)}")


def extract_name_from_resume(text: str) -> str:
    """
    Extract name from resume text - aggressive extraction.
    Uses multiple simple strategies to find any reasonable name.
    """
    if not text or len(text.strip()) < 5:
        return "Unknown Candidate"
    
    lines = [l.strip() for l in text.split("\n") if l.strip() and len(l.strip()) > 1]
    
    # Keywords that indicate this is a section header, not a name
    skip_keywords = [
        "email", "phone", "linkedin", "resume", "cv", "curriculum", 
        "vitae", "objective", "summary", "professional", "experience",
        "education", "skills", "certifications", "projects", "references",
        "address", "website", "github", "portfolio", "page", "period",
        "date", "company", "school", "university", "institute", "technical",
        "core competencies", "areas of expertise", "about", "profile",
        "http", "www", "@gmail", "@yahoo", "@hotmail"
    ]
    
    for line in lines[:20]:
        lower_line = line.lower()
        
        # Skip if contains section header keywords
        if any(skip in lower_line for skip in skip_keywords):
            continue
        
        # Skip if looks like a bullet point
        if line.startswith(("•", "-", "*", "◦")):
            continue
        
        # Skip if has lots of special characters or numbers
        special_count = sum(1 for c in line if c.isdigit() or c in "()[]{}.,;:")
        if special_count > len(line) * 0.3:  # More than 30% special chars
            continue
        
        words = line.split()
        word_count = len(words)
        
        # Skip if too many words (probably not a name)
        if word_count > 6:
            continue
        
        # Skip if too short
        if len(line) < 3:
            continue
        
        # Check if line looks like a name
        # Must have at least one capitalized word
        capitalized_words = [w for w in words if w and w[0].isupper()]
        
        if len(capitalized_words) == 0:
            continue
        
        # Avoid obvious non-names
        if words[0].lower() in ["mr", "ms", "mrs", "dr", "prof", "sir", "madam", "mr.", "ms.", "mrs."]:
            continue
        
        # If line is 1-5 words and starts with capital letter, likely a name
        if 1 <= word_count <= 5 and line[0].isupper():
            return line
    
    return "Unknown Candidate"
